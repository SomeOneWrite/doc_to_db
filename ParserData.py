import os
import sqlite3

from docx import Document

from Parsers.ParseMachines import ParseMachines
from Parsers.ParseMaterials import ParseMaterials
from Parsers.ParseTransports import ParseTransports
from Parsers.ParseWorkers import ParseWorkers

db_name = 'omsk.odb'

conn = sqlite3.connect(db_name)

query = conn.cursor()

class ParserData:

    def __init__(self, query):
        self.query = query
        pass

    def parse_table(self, table):
        rows = table.rows
        for row in rows:
            try:
                if len(row.cells) < 2:
                    continue
            except Exception as e:
                continue
            # ParseMachines(query).run(row)
            # ParseMaterials(query).run(row)

    def parse_machines(self, filename):
        doc = Document(filename)
        self.filename = filename
        tables = doc.tables
        print('Всего таблиц: ' + str(len(tables)))
        for table in range(0, len(tables)):
            print('Process {} table of {}'.format(table, len(tables)))
            rows = tables[table].rows
            for row in rows:
                try:
                    if len(row.cells) < 2:
                        continue
                except Exception as e:
                    continue
                ParseMachines(query).run(row)

    def parse_materials(self, filename):
        doc = Document(filename)
        self.filename = filename
        tables = doc.tables
        print('Всего таблиц: ' + str(len(tables)))
        for table in range(0, len(tables)):
            print('Process {} table of {}'.format(table, len(tables)))
            rows = tables[table].rows
            for row in rows:
                try:
                    if not row:
                        continue
                    if len(row.cells) < 2:
                        continue
                except Exception as e:
                    continue
                ParseMaterials(query).run(row)

    def parse_materials_text(self, filename):
        doc = Document(filename)
        self.filename = filename
        ParseMaterials(query).run_text(doc)

    def parse_file(self, filename):
        doc = Document(filename)
        self.filename = filename
        tables = doc.tables
        print('Всего таблиц: ' + str(len(tables)))
        for table in range(0, len(tables)):
            print('Process {} table of {}'.format(table, len(tables)))
            self.parse_table(tables[table])

    def parse_transports(self, filename):
        doc = Document(filename)
        self.filename = filename
        tables = doc.tables
        print('Всего таблиц: ' + str(len(tables)))
        for table in range(0, len(tables)):
            print('Process {} table of {}'.format(table, len(tables)))
            ParseTransports(query).run(tables[table])

    def parse_workers(self, filename):
        doc = Document(filename)
        self.filename = filename
        tables = doc.tables
        print('Всего таблиц: ' + str(len(tables)))
        for table in range(0, len(tables)):
            print('Process {} table of {}'.format(table, len(tables)))
            ParseWorkers(query).run(tables[table])


def search_file(filename):
    doc = Document(filename)
    print('parse file {}'.format(filename))
    tables = doc.tables
    db_result = query.execute('select * from materials where price isnull').fetchall()
    for table in tables:
        for row in table.rows:
            try:
                if not row:
                    continue
                if len(row.cells) < 2:
                    continue
            except Exception as e:
                continue
            for cell in row.cells:
                for item in db_result:
                    if cell.text.find(item[0]) != -1:
                        print(filename)



def build_dir(root):
    for file in os.listdir(root):
        if os.path.isdir(os.path.join(root, file)):
            build_dir(os.path.join(root, file))
            continue
        if file.endswith(".docx"):
            if file.startswith('~'): continue
            # print('Parse file with name: {}'.format(file))
            # if file.endswith("transports.docx"):
            # print('Parse transports')
            # ParserData(query).parse_transports(os.path.join(root, file))
            # # elif file.endswith(("workers.docx")):
            # print('Parse workers')
            # ParserData(query).parse_workers(os.path.join(root, file))
            # # elif file.endswith(("machines.docx")):
            # print('Parse machines')
            # ParserData(query).parse_machines(os.path.join(root, file))
            # # elif file.endswith(("materials.docx")):
            # print('Parse materials')
            search_file(os.path.join(root, file))
            # ParserData(query).parse_materials_text(os.path.join(root, file))


build_dir(r'C:\Users\Rinat\Desktop\new_omsk\каталоги')
conn.commit()
